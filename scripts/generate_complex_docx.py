"""
Generate a long, structurally complex DOCX: headings, lists, tables, bold/italic/underline,
alignment, and multiple embedded PNG figures (Pillow). For EC2/Lambda conversion stress tests.

  pip install -r scripts/requirements.txt
  python scripts/generate_complex_docx.py

Updates manifest entry id=docx_complex_stress_001 (creates it if missing).
"""
from __future__ import annotations

import hashlib
import json
import math
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw

REPO_ROOT = Path(__file__).resolve().parents[1]
RAW = REPO_ROOT / "datasets" / "raw"
MANIFEST = REPO_ROOT / "datasets" / "manifest.json"

DOC_ID = "docx_complex_stress_001"
OUT_NAME = f"{DOC_ID}.docx"


def sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _png_gradient(path: Path, w: int, h: int) -> None:
    img = Image.new("RGB", (w, h))
    px = img.load()
    for y in range(h):
        for x in range(w):
            r = int(40 + 180 * (x / w))
            g = int(60 + 120 * (y / h))
            b = int(200 - 100 * (x + y) / (w + h))
            px[x, y] = (r % 256, g % 256, b % 256)
    img.save(path, "PNG", optimize=True)


def _png_bar_chart(path: Path) -> None:
    w, h = 1200, 800
    img = Image.new("RGB", (w, h), (248, 249, 250))
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, w - 1, h - 1), outline=(80, 80, 100), width=2)
    vals = [120, 200, 90, 240, 160, 210, 95]
    n = len(vals)
    bw = (w - 200) // n
    m = max(vals)
    for i, v in enumerate(vals):
        x0 = 100 + i * (bw + 10)
        bar_h = int((h - 200) * (v / m))
        y0 = h - 120 - bar_h
        d.rectangle(
            (x0, y0, x0 + bw - 5, h - 120),
            fill=(70 + i * 25, 100, 180 - i * 15),
        )
    d.text((w // 2 - 120, 20), "Synthetic throughput (arbitrary units)", fill=(20, 20, 40))
    for i, v in enumerate(vals):
        d.text((100 + i * (bw + 10) + 5, h - 100), str(v), fill=(0, 0, 0))
    img.save(path, "PNG", optimize=True)


def _png_line_plot(path: Path) -> None:
    w, h = 1100, 700
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    pts = [
        (int(80 + 900 * t / 40), int(100 + 250 * (math.sin(t * 0.4) * 0.5 + 0.5) + 50 * math.sin(t * 0.2)))
        for t in range(0, 41, 1)
    ]
    for i in range(len(pts) - 1):
        d.line([pts[i], pts[i + 1]], fill=(200, 50, 50), width=3)
    d.ellipse((pts[0][0] - 5, pts[0][1] - 5, pts[0][0] + 5, pts[0][1] + 5), fill=(0, 80, 200))
    d.ellipse(
        (pts[-1][0] - 5, pts[-1][1] - 5, pts[-1][0] + 5, pts[-1][1] + 5),
        fill=(0, 150, 80),
    )
    d.text((60, 30), "Latency envelope (synthetic series)", fill=(0, 0, 0))
    img.save(path, "PNG", optimize=True)


def _png_heatmap(path: Path) -> None:
    w, h = 900, 700
    img = Image.new("RGB", (w, h))
    px = img.load()
    for y in range(h):
        for x in range(w):
            v = (math.sin(x * 0.02) * math.cos(y * 0.02) + 1) * 0.5
            c = int(40 + v * 200)
            px[x, y] = (c, 80, 255 - c // 2)
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, w - 1, h - 1), outline=(0, 0, 0), width=1)
    d.text((20, 20), "Heatmap: synthetic resource utilization", fill=(255, 255, 255))
    img.save(path, "PNG", optimize=True)


def _png_text_diagram(path: Path) -> None:
    w, h = 1000, 500
    img = Image.new("RGB", (w, h), (235, 240, 255))
    d = ImageDraw.Draw(img)
    d.rectangle((50, 50, 450, 200), fill=(200, 230, 255), outline=(50, 50, 120), width=2)
    d.rectangle((500, 50, 950, 200), fill=(255, 230, 200), outline=(120, 80, 50), width=2)
    d.polygon([(500, 350), (650, 250), (800, 350)], fill=(100, 180, 100), outline=(0, 60, 0))
    d.text((200, 95), "EC2 (steady)", fill=(0, 0, 100))
    d.text((700, 95), "Lambda (scale)", fill=(100, 40, 0))
    d.text((580, 300), "Client", fill=(0, 60, 0))
    d.line([(450, 125), (500, 125)], fill=(0, 0, 0), width=2)
    img.save(path, "PNG", optimize=True)


def set_cell_shade(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_rich_paragraph(doc: Document, text: str, *, bold=False, italic=False, underline=False, size_pt=11) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.size = Pt(size_pt)
    r.font.name = "Calibri"
    if r._element.rPr is not None and r._element.rPr.rFonts is not None:
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def build_docx(path: Path, tmpdir: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    p_grad = tmpdir / "fig_gradient.png"
    p_bar = tmpdir / "fig_bar.png"
    p_line = tmpdir / "fig_line.png"
    p_heat = tmpdir / "fig_heat.png"
    p_diagram = tmpdir / "fig_diagram.png"

    _png_gradient(p_grad, 1400, 1000)
    _png_bar_chart(p_bar)
    _png_line_plot(p_line)
    _png_heatmap(p_heat)
    _png_text_diagram(p_diagram)

    doc = Document()
    t = doc.add_paragraph("Group 259 — Complex synthetic workload for document conversion")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("CS5296 · Headings, tables, lists, and embedded figures")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_heading("1 Introduction", level=1)
    for i in range(8):
        _add_rich_paragraph(
            doc,
            f"Section 1 — Paragraph {i + 1}: Lorem-style filler for layout stress. "
            f"Conversion pipelines must handle mixed fonts, simulated emphasis, and long lines without corruption.",
            size_pt=11,
        )
    p = doc.add_paragraph()
    r0 = p.add_run("This sentence has ")
    r1 = p.add_run("bold")
    r1.bold = True
    r2 = p.add_run(", ")
    r3 = p.add_run("italic")
    r3.italic = True
    r4 = p.add_run(", ")
    r5 = p.add_run("underline")
    r5.underline = True
    r6 = p.add_run(", and combined ")
    r7 = p.add_run("bold+italic")
    r7.bold = True
    r7.italic = True
    r8 = p.add_run(" within one paragraph.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("1.1 Related patterns", level=2)
    doc.add_paragraph("Unordered list (bullets):", style="List Bullet")
    for item in [
        "Cold start vs always-on",
        "API Gateway 29s limit",
        "S3 as single source of truth",
        "CloudWatch Init Duration",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("Numbered list:", style="List Number")
    for n, item in enumerate(["Design", "Measure", "Compare", "Report"], 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"Step {n}: {item}")

    run_fig1 = doc.add_paragraph().add_run()
    run_fig1.add_picture(str(p_grad), width=Inches(5.5))
    cap1 = doc.add_paragraph("Figure: full-width color gradient (high resolution).")
    cap1.runs[0].italic = True
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading("2 Methodology", level=1)
    doc.add_heading("2.1 Workload mix", level=2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Role"
    hdr[2].text = "Bytes (approx)"
    hdr[3].text = "Notes"
    for c in hdr:
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
    data_rows = [
        ("DOCX", "input", "36–400k+", "LibreOffice path"),
        ("PDF", "input", "0.4–15k", "poppler path"),
        ("PNG", "embedded", "5 figures", "Pillow-generated"),
    ]
    for a, b, c, d in data_rows:
        row = tbl.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = a, b, c, d
    set_cell_shade(hdr[0], "D9E1F2")
    set_cell_shade(hdr[1], "D9E1F2")
    set_cell_shade(hdr[2], "D9E1F2")
    set_cell_shade(hdr[3], "D9E1F2")

    p_cap = doc.add_paragraph("Table: synthetic workload components.")
    p_cap.runs[0].italic = True
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_fig2 = doc.add_paragraph().add_run()
    run_fig2.add_picture(str(p_bar), width=Inches(5.0))
    c2 = doc.add_paragraph("Figure: bar chart (raster).")
    c2.runs[0].italic = True
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_fig3 = doc.add_paragraph().add_run()
    run_fig3.add_picture(str(p_line), width=Inches(5.2))
    c3 = doc.add_paragraph("Figure: line series.")
    c3.runs[0].italic = True
    c3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading("3 Experiments and observations", level=1)

    tbl2 = doc.add_table(rows=4, cols=3)
    tbl2.style = "Table Grid"
    for j, h in enumerate(["Metric", "EC2 (local path)", "Lambda (FURL)"]):
        tbl2.rows[0].cells[j].text = h
    metrics = [("E2E latency", "curl time_total", "Invoke-RestMethod"), ("Stability", "t3.micro 1:1", "concurrent requests")]
    for r, (m, e, l) in enumerate(metrics, 1):
        tbl2.rows[r].cells[0].text = m
        tbl2.rows[r].cells[1].text = e
        tbl2.rows[r].cells[2].text = l
    doc.add_paragraph()

    for i in range(60):
        _add_rich_paragraph(
            doc,
            f"Long-form paragraph {i + 1}: cloud benchmarking requires controlled inputs, "
            f"repeated trials, and explicit definitions of end-to-end timing. "
            f"Index i={i} embeds no secrets — purely padding for CPU/memory during PDF export.",
        )

    run_fig4 = doc.add_paragraph().add_run()
    run_fig4.add_picture(str(p_heat), width=Inches(4.5))
    c4 = doc.add_paragraph("Figure: pseudo heatmap.")
    c4.runs[0].italic = True
    c4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading("4 Architecture sketch", level=1)
    run_fig5 = doc.add_paragraph().add_run()
    run_fig5.add_picture(str(p_diagram), width=Inches(5.8))
    c5 = doc.add_paragraph("Figure: block diagram (labels are illustrative).")
    c5.runs[0].italic = True
    c5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.1 Concluding remarks (in-doc)", level=2)
    quote = doc.add_paragraph("This report section is part of a synthetic document used only to stress the conversion service.")
    quote.paragraph_format.left_indent = Inches(0.5)
    quote.paragraph_format.right_indent = Inches(0.5)
    quote.runs[0].italic = True
    quote.runs[0].font.size = Pt(10)

    for i in range(50):
        p = doc.add_paragraph()
        r = p.add_run(f"Appendix A.{i+1} — Additional narrative block. " * 2)
        r.font.size = Pt(10)
        if i % 5 == 0:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if i % 7 == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("— End of complex synthetic document —").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(path))


def upsert_manifest(sha: str, nbytes: int) -> None:
    data: list[dict] = json.loads(MANIFEST.read_text(encoding="utf-8"))
    new_row = {
        "id": DOC_ID,
        "role": "stress",
        "format": "docx",
        "source": f"file:datasets/raw/{OUT_NAME}",
        "license": "synthetic",
        "sha256": sha,
        "bytes": nbytes,
        "pages": None,
        "notes": "Pillow figures + multi-level headings + tables + lists; stress conversion",
        "synthetic": True,
    }
    for i, row in enumerate(data):
        if row.get("id") == DOC_ID:
            data[i] = new_row
            break
    else:
        data.append(new_row)
    MANIFEST.write_text(json.dumps(data, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def main() -> None:
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        out = RAW / OUT_NAME
        build_docx(out, tmp)
    sha = sha256_file(out)
    upsert_manifest(sha, out.stat().st_size)
    print("Wrote", out)
    print("bytes =", out.stat().st_size, "sha256 =", sha)
    print("Updated", MANIFEST, "for", DOC_ID)


if __name__ == "__main__":
    main()
