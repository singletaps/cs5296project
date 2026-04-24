"""
Build CS5296 final report DOCX from 草稿.md.
Format: A4, 12pt Times New Roman, single line spacing (default body & tables).
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "草稿.md"
OUT_PATH = ROOT / "Group_259_report.docx"

INLINE_RE = re.compile(r"\*\*([^*]+)\*\*|\*([^*]+)\*")


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except Exception:
        pass

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_single_spacing(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def run_font_times(run, size_pt: int = 12, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except Exception:
        pass


def add_inline_runs(paragraph, text: str, size_pt: int = 12) -> None:
    if not text:
        return
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos : m.start()])
            run_font_times(r, size_pt)
        g1, g2 = m.group(1), m.group(2)
        if g1 is not None:
            r = paragraph.add_run(g1)
            run_font_times(r, size_pt, bold=True)
        else:
            r = paragraph.add_run(g2)
            run_font_times(r, size_pt, italic=True)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        run_font_times(r, size_pt)


def add_paragraph_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_single_spacing(p)
    add_inline_runs(p, text, 12)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_single_spacing(p)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_inline_runs(p, text, 14)
    for r in p.runs:
        r.bold = True


def add_heading_level(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    set_single_spacing(p)
    add_inline_runs(p, text, 12)
    for r in p.runs:
        r.bold = True
    # Spacing similar to Word built-in headings
    p.paragraph_format.space_before = Pt(12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(4)


def is_table_sep(line: str) -> bool:
    t = line.strip()
    if not t.startswith("|"):
        return False
    return bool(re.search(r"---", t))


def split_md_row(line: str) -> list[str]:
    """Parse a markdown table row (leading pipe ... trailing pipe) into cells."""
    raw = line.strip()
    if not raw.startswith("|") or not raw.endswith("|"):
        return [c.strip() for c in raw.split("|") if c.strip() != ""]
    inner = raw[1:-1]
    return [c.strip() for c in inner.split("|")]


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i]
        if not line.strip().startswith("|"):
            break
        if is_table_sep(line):
            i += 1
            continue
        rows.append(split_md_row(line))
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_col = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_col)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(n_col):
            cell = tbl.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            p.text = ""
            set_single_spacing(p)
            val = row[ci] if ci < len(row) else ""
            add_inline_runs(p, val, 11 if n_col >= 6 else 12)
    doc.add_paragraph()


def build_doc(md: str) -> Document:
    doc = Document()
    set_document_defaults(doc)

    lines = md.splitlines()
    i = 0
    title_used = False

    while i < len(lines):
        raw = lines[i]
        s = raw.strip()

        if s == "":
            i += 1
            continue
        if s == "---":
            i += 1
            continue

        # First level-1 heading: document title (centered), not "##"
        if s.startswith("# ") and not s.startswith("##") and not title_used:
            add_title(doc, s[2:].strip())
            title_used = True
            i += 1
            continue

        if s.startswith("## "):
            add_heading_level(doc, s[3:].strip(), 2)
            i += 1
            continue

        if s.startswith("### "):
            add_heading_level(doc, s[4:].strip(), 3)
            i += 1
            continue

        if s.startswith("|"):
            rows, j = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
            i = j
            continue

        add_paragraph_body(doc, s)
        i += 1

    return doc


def main() -> int:
    md_path = Path(sys.argv[1]) if len(sys.argv) > 1 else MD_PATH
    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else OUT_PATH
    if not md_path.is_file():
        print(f"Missing file: {md_path}", file=sys.stderr)
        return 1
    text = md_path.read_text(encoding="utf-8")
    doc = build_doc(text)
    out_path = out_path.resolve()
    try:
        doc.save(str(out_path))
    except PermissionError:
        alt = out_path.with_name(out_path.stem + "_new" + out_path.suffix)
        doc.save(str(alt))
        print(f"Wrote {alt} (default path was locked; close Word and re-run to overwrite)")
        return 0
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
